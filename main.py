from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def modify_word_file(input_file, output_file):
    doc = Document(input_file)
    new_font_size = Pt(12)
    new_font_style = 'Arial'

    # Iterate through paragraphs and modify font size and style
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = new_font_size
            run.font.name = new_font_style

    # Iterate through tables and modify font size and style
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = new_font_size
                        run.font.name = new_font_style

    # Save the modified document to a new file
    doc.save(output_file)

if __name__ == "__main__":
    input_file_path = '/home/ahmed/CV.docx'
    output_file_path = '/home/ahmed/CV2.docx'

    modify_word_file(input_file_path, output_file_path)
